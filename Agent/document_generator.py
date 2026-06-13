"""
公文生成模块 - 支持带格式的公文生成

归档说明：本模块是早期 DashScope/LangChain 生成器示例，不属于当前
Flask/LangGraph 主服务。需要运行时请安装 requirements-examples.txt。

使用流程：
1. 用户输入需求
2. 检索知识库获取模板和规范
3. 构建 Prompt（含格式要求）
4. 调用 Qwen API 生成
5. 应用格式导出为 docx
"""

import os
from typing import List, Dict, Optional
from dataclasses import dataclass
import json

try:
    from langchain_community.vectorstores import FAISS
    from langchain_community.embeddings import DashScopeEmbeddings
    from dashscope import Generation
except ModuleNotFoundError:
    FAISS = DashScopeEmbeddings = Generation = None

# 文档处理
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# 格式处理
from document_parser import chinese_size_to_pt, FormatFingerprint


@dataclass
class GenerationRequest:
    """公文生成请求"""
    requirement: str           # 用户需求
    document_type: str = ""    # 公文类型（通知、请示等）
    context: str = ""         # 额外上下文


@dataclass
class GenerationResult:
    """公文生成结果"""
    title: str                # 标题
    content: str              # 正文内容
    format_spec: dict         # 格式规范
    metadata: dict            # 元数据


class DocumentGenerator:
    """公文生成器"""

    def __init__(self, knowledge_base_path: str):
        """
        初始化公文生成器

        Args:
            knowledge_base_path: 知识库路径
        """
        if FAISS is None or DashScopeEmbeddings is None or Generation is None:
            raise RuntimeError("document_generator 已归档；请先安装 Agent/requirements-examples.txt 中的可选依赖")

        # 加载 FAISS 索引
        self.embeddings = DashScopeEmbeddings(
            model="text-embedding-v2",
            dashscope_api_key=os.getenv("DASHSCOPE_API_KEY")
        )

        self.vectorstore = FAISS.load_local(
            knowledge_base_path,
            self.embeddings
        )

        # 加载配置
        config_path = os.path.join(knowledge_base_path, "config.json")
        if os.path.exists(config_path):
            with open(config_path, "r", encoding="utf-8") as f:
                self.config = json.load(f)
        else:
            self.config = {}

    def retrieve_context(self, query: str, k: int = 5) -> List[Dict]:
        """
        检索相关上下文

        Args:
            query: 查询内容
            k: 返回数量

        Returns:
            检索结果列表
        """
        results = self.vectorstore.similarity_search(query, k=k)

        context = []
        for doc in results:
            item = {
                "content": doc.page_content,
                "metadata": doc.metadata
            }

            # 如果有格式信息，也提取出来
            if "format" in doc.metadata:
                item["format"] = doc.metadata["format"]

            context.append(item)

        return context

    def extract_format_rules(self, context: List[Dict]) -> dict:
        """
        从上下文中提取格式规则

        Args:
            context: 检索到的上下文

        Returns:
            格式规则字典
        """
        # 默认格式规则（国家标准公文格式）
        default_rules = {
            "标题": {
                "font": "黑体",
                "size": "二号",
                "bold": True,
                "alignment": "居中"
            },
            "正文": {
                "font": "仿宋_GB2312",
                "size": "三号",
                "bold": False,
                "alignment": "左对齐",
                "line_spacing": "28磅",
                "first_line_indent": 2
            },
            "落款": {
                "font": "仿宋_GB2312",
                "size": "三号",
                "bold": False,
                "alignment": "居右"
            }
        }

        # 从上下文中查找格式规范
        for item in context:
            if "format" in item:
                fmt = item["format"]
                # 根据内容判断是标题、正文还是落款
                content = item["content"]

                if any(keyword in content for keyword in ["标题", "题目", "关于"]):
                    default_rules["标题"].update(fmt)
                elif any(keyword in content for keyword in ["落款", "署名", "日期", "部门"]):
                    default_rules["落款"].update(fmt)
                else:
                    default_rules["正文"].update(fmt)

        return default_rules

    def build_prompt(self, request: GenerationRequest, context: List[Dict]) -> str:
        """
        构建 Prompt

        Args:
            request: 生成请求
            context: 检索到的上下文

        Returns:
            完整的 Prompt
        """
        # 提取格式规则
        format_rules = self.extract_format_rules(context)

        # 构建格式描述
        format_desc = self._format_rules_to_text(format_rules)

        # 构建上下文描述
        context_desc = self._build_context_desc(context)

        # 构建 Prompt
        prompt = f"""你是一位公文写作专家，擅长根据用户需求生成符合规范的公文。

## 用户需求
{request.requirement}

## 格式规范
{format_desc}

## 参考范文
{context_desc}

## 生成要求
1. 严格按照上述格式规范生成公文
2. 内容要准确、完整、正式
3. 语言要简洁、明了
4. 标题居中，正文首行缩进2字符
5. 落款居右，包括部门名称和日期

## 输出格式
请直接输出公文内容，不需要额外说明。
"""
        return prompt

    def _format_rules_to_text(self, rules: dict) -> str:
        """将格式规则转换为文本描述"""
        lines = []

        for section, fmt in rules.items():
            line = f"{section}："
            specs = []

            if fmt.get("font"):
                specs.append(f"字体：{fmt['font']}")
            if fmt.get("size"):
                specs.append(f"字号：{fmt['size']}")
            if fmt.get("bold"):
                specs.append("加粗")
            if fmt.get("alignment"):
                specs.append(f"对齐：{fmt['alignment']}")
            if fmt.get("line_spacing"):
                specs.append(f"行距：{fmt['line_spacing']}")
            if fmt.get("first_line_indent"):
                specs.append(f"首行缩进：{fmt['first_line_indent']}字符")

            line += "，".join(specs) if specs else "默认格式"
            lines.append(line)

        return "\n".join(lines)

    def _build_context_desc(self, context: List[Dict]) -> str:
        """构建上下文描述"""
        if not context:
            return "无参考范文"

        lines = []
        for i, item in enumerate(context[:3]):  # 最多3个参考
            content = item["content"]
            filename = item["metadata"].get("filename", "未知")
            lines.append(f"【参考{i+1}】{filename}：\n{content[:200]}...")

        return "\n\n".join(lines)

    def generate(self, request: GenerationRequest) -> Optional[GenerationResult]:
        """
        生成公文

        Args:
            request: 生成请求

        Returns:
            生成结果
        """
        # 1. 检索相关上下文
        context = self.retrieve_context(request.requirement, k=5)

        # 2. 构建 Prompt
        prompt = self.build_prompt(request, context)

        # 3. 调用 Qwen API
        response = Generation.call(
            model="qwen-max",
            prompt=prompt,
            api_key=os.getenv("DASHSCOPE_API_KEY")
        )

        if response.status_code != 200:
            print(f"❌ API 调用失败：{response.message}")
            return None

        # 4. 解析结果
        generated_text = response.output.text

        # 5. 提取标题和正文
        result = self._parse_generated_text(generated_text)

        # 6. 提取格式规范
        format_spec = self.extract_format_rules(context)

        return GenerationResult(
            title=result.get("title", ""),
            content=result.get("content", ""),
            format_spec=format_spec,
            metadata={
                "document_type": request.document_type,
                "context_count": len(context)
            }
        )

    def _parse_generated_text(self, text: str) -> dict:
        """解析生成的文本，提取标题和正文"""
        lines = text.strip().split("\n")

        result = {
            "title": "",
            "content": text
        }

        # 简单策略：第一行非空且短的是标题
        for i, line in enumerate(lines):
            if line.strip() and len(line.strip()) < 30 and i < 5:
                result["title"] = line.strip()
                result["content"] = "\n".join(lines[i+1:])
                break

        return result

    def export_to_docx(self, result: GenerationResult, output_path: str):
        """
        导出为 docx 文件

        Args:
            result: 生成结果
            output_path: 输出路径
        """
        doc = Document()

        # 设置默认样式
        style = doc.styles["Normal"]
        style.font.name = "仿宋_GB2312"
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

        # 1. 标题
        if result.title:
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(result.title)
            title_run.bold = True
            title_run.font.size = Pt(chinese_size_to_pt("二号"))
            title_run.font.name = "黑体"
            title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 2. 正文
        content_lines = result.content.strip().split("\n")
        for line in content_lines:
            if not line.strip():
                continue

            para = doc.add_paragraph()
            para_run = para.add_run(line)

            # 应用正文格式
            para_run.font.size = Pt(chinese_size_to_pt("三号"))
            para_run.font.name = "仿宋_GB2312"
            para_run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

            # 首行缩进
            para.paragraph_format.first_line_indent = Cm(0.74)  # 2字符 ≈ 0.74cm

            # 行距
            para.paragraph_format.line_spacing = Pt(28)

        # 3. 落款（最后几行）
        # 检测落款部分并右对齐
        paras = doc.paragraphs
        if len(paras) > 3:
            # 最后3个段落作为落款
            for i in range(max(0, len(paras)-3), len(paras)):
                paras[i].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # 保存
        doc.save(output_path)
        print(f"✅ 公文已导出：{output_path}")


def generate_document(requirement: str, knowledge_base_path: str, output_path: str):
    """
    便捷函数：生成公文并导出

    Args:
        requirement: 用户需求
        knowledge_base_path: 知识库路径
        output_path: 输出路径
    """
    generator = DocumentGenerator(knowledge_base_path)

    request = GenerationRequest(requirement=requirement)
    result = generator.generate(request)

    if result:
        generator.export_to_docx(result, output_path)
        return result
    else:
        print("❌ 生成失败")
        return None


if __name__ == "__main__":
    # 测试
    import sys

    if len(sys.argv) > 2:
        requirement = sys.argv[1]
        output_path = sys.argv[2]
        kb_path = "./knowledge_base"

        print("=" * 60)
        print("测试公文生成")
        print("=" * 60)

        result = generate_document(requirement, kb_path, output_path)

        if result:
            print(f"\n📄 生成结果：")
            print(f"标题：{result.title}")
            print(f"正文：{result.content[:200]}...")
            print(f"\n格式规范：")
            for section, fmt in result.format_spec.items():
                print(f"  {section}：{fmt}")
