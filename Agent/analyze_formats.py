#!/usr/bin/env python3
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import json

docs_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), '知识库')

def analyze_run(run):
    info = {
        'text': run.text[:30] if run.text else '',
        'font_name': None,
        'font_size': None,
        'bold': None,
        'italic': None,
        'underline': None,
    }
    try:
        if run.font and run.font.name:
            info['font_name'] = run.font.name
        if run.font and run.font.size:
            info['font_size'] = round(run.font.size.pt, 1)
        if run.font and run.font.bold is not None:
            info['bold'] = run.font.bold
        if run.font and run.font.italic is not None:
            info['italic'] = run.font.italic
        if run.font and run.font.underline is not None:
            info['underline'] = run.font.underline
    except:
        pass
    return info

results = {}
files = [f for f in os.listdir(docs_dir) if f.endswith('.docx')]

for fname in files:
    path = os.path.join(docs_dir, fname)
    try:
        doc = docx.Document(path)
        all_runs = []
        all_formats = []

        for para in doc.paragraphs:
            try:
                if para.paragraph_format:
                    pf = para.paragraph_format
                    fmt = {
                        'alignment': str(pf.alignment) if pf.alignment else None,
                        'line_spacing': pf.line_spacing if pf.line_spacing else None,
                        'first_line_indent': pf.first_line_indent if pf.first_line_indent else None,
                    }
                    all_formats.append(fmt)
            except:
                pass

            for run in para.runs:
                run_info = analyze_run(run)
                if run_info['text'].strip():
                    all_runs.append(run_info)

        unique_font_combos = {}
        for r in all_runs:
            key = (r['font_name'], r['font_size'], r['bold'])
            if key not in unique_font_combos:
                unique_font_combos[key] = 0
            unique_font_combos[key] += 1

        results[fname] = {
            'total_paragraphs': len(doc.paragraphs),
            'total_runs': len(all_runs),
            'unique_formats_count': len(unique_font_combos),
            'format_distribution': {str(k): v for k, v in unique_font_combos.items()},
            'paragraph_formats_sample': all_formats[:10]
        }
    except Exception as e:
        results[fname] = {'error': str(e)}

print(json.dumps(results, ensure_ascii=False, indent=2))