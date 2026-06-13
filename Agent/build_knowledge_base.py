#!/usr/bin/env python3
import os
import json
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

docs_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), '知识库')

def extract_text_from_docx(path):
    doc = docx.Document(path)
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            para_data = {
                'text': para.text,
                'alignment': str(para.alignment) if para.alignment else 'LEFT',
                'runs_info': []
            }
            for run in para.runs:
                try:
                    para_data['runs_info'].append({
                        'font': run.font.name if run.font and run.font.name else '默认',
                        'size': round(run.font.size.pt, 1) if run.font and run.font.size else '默认',
                        'bold': run.font.bold if run.font else False,
                        'text_preview': run.text[:50]
                    })
                except:
                    pass
            paragraphs.append(para_data)
    return paragraphs

def analyze_single_doc(path, filename):
    result = {
        'filename': filename,
        'text_samples': [],
        'format_patterns': set()
    }

    try:
        doc = docx.Document(path)

        for para in doc.paragraphs[:50]:
            if not para.text.strip():
                continue

            text = para.text.strip()
            if len(text) < 10:
                continue

            para_info = {
                'text': text[:500],
                'alignment': None,
                'styles': []
            }

            if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                para_info['alignment'] = '居中'
            elif para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                para_info['alignment'] = '两端对齐'
            elif para.alignment == WD_ALIGN_PARAGRAPH.LEFT:
                para_info['alignment'] = '左对齐'

            for run in para.runs:
                try:
                    if run.font:
                        font_info = {
                            'font': run.font.name or '默认',
                            'size': round(run.font.size.pt, 1) if run.font.size else '默认',
                            'bold': run.font.bold or False
                        }
                        para_info['styles'].append(font_info)

                        key = f"{font_info['font']}_{font_info['size']}号_{'加粗' if font_info['bold'] else '常规'}"
                        result['format_patterns'].add(key)
                except:
                    pass

            result['text_samples'].append(para_info)

    except Exception as e:
        result['error'] = str(e)

    result['format_patterns'] = list(result['format_patterns'])
    return result

def get_all_docs_analysis():
    all_results = []
    files = [f for f in os.listdir(docs_dir) if f.endswith('.docx')]

    for fname in files:
        path = os.path.join(docs_dir, fname)
        print(f"分析中: {fname}")
        result = analyze_single_doc(path, fname)
        all_results.append(result)

    return all_results

if __name__ == '__main__':
    print("=" * 60)
    print("开始分析知识库文档格式...")
    print("=" * 60)

    results = get_all_docs_analysis()

    print("\n" + "=" * 60)
    print("格式统计汇总")
    print("=" * 60)

    all_patterns = {}
    for r in results:
        for p in r.get('format_patterns', []):
            if p not in all_patterns:
                all_patterns[p] = 0
            all_patterns[p] += 1

    print("\n使用的字体格式（按频率）:")
    for pattern, count in sorted(all_patterns.items(), key=lambda x: -x[1]):
        print(f"  {pattern}: {count}个文档使用")

    print("\n" + "=" * 60)
    print("样本文本内容预览（前3个文档）")
    print("=" * 60)

    for r in results[:3]:
        print(f"\n【{r['filename']}】")
        for sample in r['text_samples'][:2]:
            print(f"  文本: {sample['text'][:100]}...")
            print(f"  样式: {sample['styles']}")

    output_path = '/Users/qfen9/Documents/code/Agent/knowledge_analysis.json'
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(results, f, ensure_ascii=False, indent=2)
    print(f"\n详细分析结果已保存到: {output_path}")