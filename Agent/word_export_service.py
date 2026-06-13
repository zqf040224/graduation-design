"""Word and template helpers for document export."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def parse_document_format(text):
    lines = text.split('\n')
    formatted_parts = []
    non_empty = [line.strip() for line in lines if line.strip()]
    if not non_empty:
        return formatted_parts

    total = len(non_empty)
    cn_digits = '一二三四五六七八九十'
    heading1 = re.compile(rf'^[{cn_digits}]{{1,3}}[、，]')
    heading2 = re.compile(rf'^（[{cn_digits}{cn_digits}]{{1,2}}）')
    heading3 = re.compile(r'^\d+[\.\)、）]')
    date_pattern = re.compile(r'\d{4}\s*年\s*\d{1,2}\s*月\s*\d{1,2}\s*日')
    closings = {
        '特此通知。', '特此报告。', '特此通告。', '特此函告。', '特此批复。',
        '特此请示。', '以上建议，请批示。', '以上报告，请审阅。',
        '妥否，请批示。', '当否，请批示。', '此复。', '此告。',
    }

    for index, line in enumerate(non_empty):
        part_type = 'body'
        if index == 0 and not heading1.match(line) and not heading2.match(line):
            part_type = 'title'
        elif any(line.startswith(prefix) for prefix in [
            '各部门', '各相关单位', '各学院', '各职能部门', '单位负责人',
            '各有关单位', '各教学科研单位', '各处室', '各直属单位',
            '各学院（系）', '各相关部门',
        ]):
            part_type = 'recipient'
        elif heading1.match(line):
            part_type = 'heading1'
        elif heading2.match(line):
            part_type = 'heading2'
        elif heading3.match(line) and len(line) < 50:
            part_type = 'heading3'
        elif line.startswith(('附件：', '附件:')):
            part_type = 'attachment'
        elif line in closings or (len(line) < 30 and any(line.endswith(closing[-2:]) for closing in closings)):
            part_type = 'closing'
        elif line.startswith(('（联系人', '联系人', '联系电话')):
            part_type = 'contact'
        elif date_pattern.search(line) and index >= total - 4:
            part_type = 'signature_date'
        elif index >= total - 5 and len(line) < 30 and any(
            keyword in line for keyword in ['大学', '学院', '单位', '办公室', '委员会', '中心']
        ):
            part_type = 'signature_unit'
        formatted_parts.append({'type': part_type, 'text': line})

    return formatted_parts


def detect_export_template(text: str, requested_template: str = "") -> str:
    requested = (requested_template or "").strip().lower()
    if requested in {"review_proposal", "proposal", "议案", "审议", "院务会议案"}:
        return "review_proposal"
    if requested and requested not in {"auto", "default", "公文"}:
        return "default"

    sample = (text or "")[:500]
    review_markers = ["议案", "审议", "院务会", "有关事宜", "此议案需决议", "以上，请审议"]
    if any(marker in sample for marker in review_markers):
        return "review_proposal"
    return "default"


def resolve_export_template(text: str, plan: dict = None, user_request: str = "") -> str:
    plan = plan or {}
    document_type = str(plan.get("document_type") or "")
    task_type = str(plan.get("task_type") or "")
    explicit_hint = document_type or task_type
    combined_text = "\n".join(part for part in [user_request or "", text or ""] if part)
    return detect_export_template(combined_text, explicit_hint or "auto")


def detect_reimbursement_template(text: str, requested_template: str = "") -> str:
    requested = (requested_template or "").strip().lower()
    aliases = {
        "travel": "travel",
        "差旅": "travel",
        "差旅费": "travel",
        "chailv": "travel",
        "meeting": "meeting",
        "会议": "meeting",
        "会议费": "meeting",
        "labor": "labor_expert",
        "expert": "labor_expert",
        "labor_expert": "labor_expert",
        "劳务": "labor_expert",
        "劳务费": "labor_expert",
        "专家": "labor_expert",
        "专家咨询费": "labor_expert",
        "other": "other",
        "其他": "other",
        "其他费用": "other",
        "其他费用报销": "other",
    }
    if requested in aliases:
        return aliases[requested]
    if requested and requested not in {"auto", "default", "xlsx", "excel"}:
        return ""

    sample = re.sub(r"\s+", "", text or "")[:800]
    if not sample:
        return ""
    if any(marker in sample for marker in ("劳务费", "劳务报销", "专家咨询费", "专家费", "专家咨询")):
        return "labor_expert"
    if any(marker in sample for marker in ("差旅费", "差旅报销", "出差报销", "外出报销")):
        return "travel"
    if any(marker in sample for marker in ("会议费", "会议报销", "办会报销")):
        return "meeting"
    if any(marker in sample for marker in ("其他费用报销", "其他报销", "其他费用")):
        return "other"
    return ""


def resolve_reimbursement_template(text: str, plan: dict = None, user_request: str = "") -> str:
    plan = plan or {}
    plan_text = "\n".join(
        str(plan.get(key) or "")
        for key in ("document_type", "task_type", "title")
    )
    combined_text = "\n".join(part for part in [user_request or "", plan_text, text or ""] if part)
    return detect_reimbursement_template(combined_text, "auto")


def reimbursement_template_path(template_key: str, *, template_dir: Path, template_files: dict[str, str]) -> Path:
    filename = template_files.get(template_key or "")
    if not filename:
        return Path()
    return template_dir / filename


def explicit_document_request(text: str, *, has_file_content: bool = False) -> bool:
    sample = re.sub(r"\s+", "", text or "")
    if not sample:
        return False
    format_markers = (
        "改为公文格式", "改成公文格式", "改成公文", "转成公文",
        "套用公文", "规范公文", "公文格式", "按公文", "正式公文",
    )
    if any(marker in sample for marker in format_markers):
        return True
    generation_verbs = ("写一份", "起草", "拟一份", "拟写", "生成", "撰写", "出一份", "帮我写")
    document_nouns = ("通知", "请示", "报告", "函", "议案", "审议", "会议纪要", "方案", "制度", "办法", "公文")
    if any(verb in sample for verb in generation_verbs) and any(noun in sample for noun in document_nouns):
        return True
    return any(marker in sample.lower() for marker in ("导出word", "导出docx", "生成word", "生成docx"))


def create_docx(text, template_type: str = "default", *, review_template_path: Path | None = None):
    if detect_export_template(text, template_type) == "review_proposal":
        return create_review_proposal_docx(text, review_template_path=review_template_path)

    doc = Document()
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(3.7)
        section.bottom_margin = Cm(3.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.6)

    font_size_map = {
        'title': 22, 'recipient': 16, 'heading1': 16, 'heading2': 16,
        'heading3': 16, 'body': 16, 'attachment': 16, 'closing': 16,
        'signature_unit': 16, 'signature_date': 16, 'contact': 16,
    }
    font_name_map = {
        'title': '方正小标宋简体', 'recipient': '仿宋_GB2312',
        'heading1': '黑体', 'heading2': '楷体', 'heading3': '仿宋_GB2312',
        'body': '仿宋_GB2312', 'attachment': '黑体', 'closing': '仿宋_GB2312',
        'signature_unit': '仿宋_GB2312', 'signature_date': '仿宋_GB2312',
        'contact': '仿宋_GB2312',
    }
    p_style_map = {
        'title': 'center', 'recipient': 'left', 'heading1': 'left',
        'heading2': 'left', 'heading3': 'left', 'body': 'justify',
        'attachment': 'left', 'closing': 'left', 'signature_unit': 'center',
        'signature_date': 'right', 'contact': 'left',
    }

    for part in parse_document_format(text):
        if part['type'] == 'empty':
            continue
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(part['text'])
        font_size = font_size_map.get(part['type'], 16)
        font_name = font_name_map.get(part['type'], '仿宋_GB2312')
        paragraph_style = p_style_map.get(part['type'], 'left')

        run.font.size = Pt(font_size)
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = Pt(28)
        if part['type'] == 'title':
            run.font.bold = True
        if paragraph_style == 'center':
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif paragraph_style == 'right':
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        elif paragraph_style == 'justify':
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            paragraph.paragraph_format.first_line_indent = Cm(0.74)
        if part['type'] == 'body' and not part['text'].startswith(('一', '（', '各', '经', '为', '现', '根', '据')):
            paragraph.paragraph_format.first_line_indent = Cm(0.74)

    return doc


def create_review_proposal_docx(text: str, *, review_template_path: Path | None = None) -> Document:
    fields = _split_review_proposal_text(text)
    doc = _review_docx_from_template(review_template_path)

    for _ in range(2):
        paragraph = doc.add_paragraph()
        _set_paragraph_base(paragraph, WD_PARAGRAPH_ALIGNMENT.CENTER)

    paragraph = doc.add_paragraph()
    _set_paragraph_base(paragraph, WD_PARAGRAPH_ALIGNMENT.CENTER)
    _set_run_font(paragraph.add_run(fields["title"]), "方正小标宋简体", 22, bold=True)

    paragraph = doc.add_paragraph()
    _set_paragraph_base(paragraph, WD_PARAGRAPH_ALIGNMENT.CENTER)

    paragraph = doc.add_paragraph()
    _set_paragraph_base(paragraph, WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
    _set_run_font(paragraph.add_run(fields["recipient"]), "仿宋", 16)

    for line in fields["body"]:
        paragraph = doc.add_paragraph()
        _set_paragraph_base(paragraph, WD_PARAGRAPH_ALIGNMENT.LEFT, first_line_indent_cm=1.13)
        _set_run_font(paragraph.add_run(line), "仿宋", 16)

    if fields["attachments"]:
        doc.add_paragraph()
        for index, line in enumerate(fields["attachments"]):
            paragraph = doc.add_paragraph()
            indent = 2.82 if index > 0 and re.match(r"^\d+[\.、]", line) else 1.13
            _set_paragraph_base(paragraph, WD_PARAGRAPH_ALIGNMENT.JUSTIFY, first_line_indent_cm=indent)
            _set_run_font(paragraph.add_run(line), "仿宋", 16)

    doc.add_paragraph()
    paragraph = doc.add_paragraph()
    _set_paragraph_base(paragraph, WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
    _set_run_font(paragraph.add_run("\u3000\u3000" + fields["decision"].lstrip("\u3000 ")), "仿宋", 16)

    doc.add_paragraph()
    doc.add_paragraph()
    paragraph = doc.add_paragraph()
    _set_paragraph_base(paragraph, WD_PARAGRAPH_ALIGNMENT.RIGHT)
    _set_run_font(paragraph.add_run(fields["signature_unit"]), "仿宋", 16)
    paragraph = doc.add_paragraph()
    _set_paragraph_base(paragraph, WD_PARAGRAPH_ALIGNMENT.RIGHT)
    _set_run_font(paragraph.add_run(fields["signature_date"]), "仿宋", 16)
    return doc


def _set_run_font(run, font_name: str, size_pt: int, bold: bool = False):
    run.font.size = Pt(size_pt)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.bold = bold


def _set_paragraph_base(paragraph, alignment=None, first_line_indent_cm=None):
    if alignment is not None:
        paragraph.alignment = alignment
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(28)
    if first_line_indent_cm is not None:
        paragraph.paragraph_format.first_line_indent = Cm(first_line_indent_cm)


def _strip_reference_block(lines: list) -> list:
    cleaned = []
    skipping = False
    for line in lines:
        if line.strip().startswith("【参考来源】"):
            skipping = True
            continue
        if skipping:
            continue
        cleaned.append(line)
    return cleaned


def _split_review_proposal_text(text: str) -> dict:
    lines = [line.strip() for line in (text or "").splitlines() if line.strip()]
    lines = _strip_reference_block(lines)
    title = "关于审议XXX的有关事宜"
    recipient = "各位领导："
    body = []
    attachments = []
    decision = ""
    signature_unit = "XXX部"
    signature_date = "XXXX年XX月XX日"

    date_re = re.compile(r"\d{4}\s*年\s*\d{1,2}\s*月\s*\d{1,2}\s*日|X{4}\s*年\s*X{1,2}\s*月\s*X{1,2}\s*日")
    for index, line in enumerate(lines):
        if index == 0 and not line.startswith(("各位领导", "附件", "此议案需决议")):
            title = line
            continue
        if line.startswith("各位领导"):
            recipient = line
            continue
        if line.startswith(("附件：", "附件:")) or (attachments and re.match(r"^\d+[\.、]", line)):
            attachments.append(line)
            continue
        if line.startswith("此议案需决议"):
            decision = line
            continue
        if date_re.search(line):
            signature_date = line
            continue
        if index >= max(len(lines) - 3, 0) and len(line) <= 30 and any(
            keyword in line for keyword in ["部", "处", "办", "中心", "学院", "单位", "委员会"]
        ):
            signature_unit = line
            continue
        body.append(line)

    if not body:
        body = [
            "XXXXXX（此处写事由）。",
            "根据《xxx制度/办法/规定/细则》第几条第几款规定：“……” ，XXXXXXX（此处写需提请决议事项）。",
            "XXXXXXX。（据实描述议案内容，说明依据、必要性和拟决议事项。）",
        ]
    if body and not any("以上，请审议" in line for line in body):
        body.append("以上，请审议。")
    if not decision:
        decision = "此议案需决议：是否通过上述事项。"

    return {
        "title": title,
        "recipient": recipient,
        "body": body,
        "attachments": attachments,
        "decision": decision,
        "signature_unit": signature_unit,
        "signature_date": signature_date,
    }


def _review_docx_from_template(review_template_path: Path | None) -> Document:
    if review_template_path and review_template_path.is_file():
        doc = Document(str(review_template_path))
        doc._body.clear_content()
        return doc
    doc = Document()
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(3.7)
        section.bottom_margin = Cm(3.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.6)
    return doc
