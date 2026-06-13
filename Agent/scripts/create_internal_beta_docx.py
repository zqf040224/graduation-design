#!/usr/bin/env python3
"""Create the internal beta usage guide DOCX."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).resolve().parents[1] / "outputs" / "Knowledge_Agent_内测使用说明.docx"


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1F2937"
MUTED = "6B7280"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
BORDER = "D9E2EC"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, width_dxa=9360, indent_dxa=120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.first_child_found_in("w:cantSplit")
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def set_cell_text(cell, text: str, bold: bool = False, color: str = INK) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(color)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    set_table_width(table)
    set_table_borders(table)
    header_row = table.rows[0]
    set_repeat_table_header(header_row)
    set_row_cant_split(header_row)
    for idx, header in enumerate(headers):
        cell = header_row.cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        cell.width = Inches(widths[idx])
        set_cell_text(cell, header, bold=True, color=DARK_BLUE)
    for row in rows:
        data_row = table.add_row()
        set_row_cant_split(data_row)
        cells = data_row.cells
        for idx, value in enumerate(row):
            cells[idx].width = Inches(widths[idx])
            set_cell_text(cells[idx], value)
    doc.add_paragraph()


def add_callout(doc: Document, title: str, body: str, fill: str = CALLOUT) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    set_table_width(table)
    set_table_borders(table, color="CBD5E1")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    run2 = p2.add_run(body)
    run2.font.size = Pt(10.5)
    run2.font.color.rgb = RGBColor.from_string(INK)
    run2.font.name = "Calibri"
    run2._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    doc.add_paragraph()


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor.from_string(INK)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_numbered(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor.from_string(INK)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header_p = section.header.paragraphs[0]
    header_p.text = "Knowledge Agent 内测使用说明"
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_p.runs[0].font.size = Pt(9)
    header_p.runs[0].font.color.rgb = RGBColor.from_string(MUTED)

    footer_p = section.footer.paragraphs[0]
    footer_p.text = "试用说明资料｜请勿公开管理员账号或测试材料"
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.runs[0].font.size = Pt(9)
    footer_p.runs[0].font.color.rgb = RGBColor.from_string(MUTED)


def build_doc() -> None:
    doc = Document()
    style_document(doc)
    add_header_footer(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run("Knowledge Agent 内测使用说明")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    sub = subtitle.add_run("适用对象：试用用户｜版本：Beta 0.1｜日期：2026年6月11日")
    sub.font.size = Pt(10.5)
    sub.font.color.rgb = RGBColor.from_string(MUTED)
    sub.font.name = "Calibri"
    sub._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    add_callout(
        doc,
        "内测目标",
        "本次内测重点验证当前系统全部主要功能：智能客服、知识库问答、Excel 表格问答、NAS/网盘标准问答、"
        "公文生成多 Agent 协作、文档编辑器、Word/Excel 导出、附件上传、知识库管理、内测反馈和日志监控。"
        "请优先使用真实办公问题测试，但不要上传高度敏感或个人隐私材料。",
    )

    doc.add_heading("一、访问入口与账号", level=1)
    add_table(
        doc,
        ["项目", "说明"],
        [
            ["访问地址", "本机测试地址：http://127.0.0.1:5003；其他同事访问需使用局域网 IP 或后续配置局域网域名。"],
            ["聊天入口", "http://127.0.0.1:5003/chat"],
            ["管理入口", "http://127.0.0.1:5003/admin，仅管理员使用。"],
            ["账号发放", "账号由管理员统一创建。请勿在群聊中公开管理员账号、密码或敏感截图。"],
        ],
        [1.45, 5.05],
    )

    doc.add_heading("二、功能指引", level=1)
    add_callout(
        doc,
        "使用原则",
        "内测时请按真实办公流程使用系统：先问答或生成，再进入文档编辑器核对，最后导出文件并提交反馈。"
        "同一功能建议至少用 2-3 种不同问法或材料测试。",
    )
    add_table(
        doc,
        ["功能", "入口", "使用指引"],
        [
            ["智能客服", "输入区下方模式按钮“智能客服”。", "适合知识库问答、制度查询、NAS/网盘地址、场地收费、材料归纳等快速问题。"],
            ["公文协作", "输入区下方模式按钮“公文协作”。", "适合通知、报告、请示、议案、对策建议等公文生成，以及连续修改、润色、压缩篇幅。"],
            ["快捷任务", "首页建议卡片：通知、汇报、建议、改写。", "点击后自动填入典型任务，可用于快速启动内测场景。"],
            ["会话管理", "左侧“新建对话”、最近会话、删除。", "用于新建测试会话、回看历史问题、删除不需要的会话。"],
            ["附件上传", "输入框左侧“添加附件”或编辑器来源区“上传”。", "支持 doc/docx/pdf/txt/md/xlsx/xls/csv；可选择加入团队知识库或仅当前对话临时使用。"],
            ["表格处理", "上传 Excel/CSV 后点击附件上的“处理”。", "可按规则筛选、排序、限制行数，并导出处理后的 Excel。"],
            ["文档编辑器", "顶部“文档编辑器”按钮或生成后自动导入右侧编辑器。", "可直接编辑正文，调整字体、字号、行距，插入来源，清空内容，查看字数和自动保存状态。"],
            ["Word 导出", "编辑器底部“Word 模板”和“导出 Word”。", "导出前选择自动识别、普通公文或院务会议案模板；导出后打开 docx 核对版式。"],
            ["Excel 导出", "编辑器底部“导出 Excel”或表格处理结果按钮。", "用于导出表格内容、报销表模板或上传表格处理结果。"],
            ["复制与反馈", "编辑器底部“复制”；顶部“反馈”。", "复制正文用于人工校对；遇到错答、漏答、导出异常、上传问题时提交内测反馈。"],
            ["后台管理", "顶部“管理”。", "管理员查看知识库文件、上传删除记录、表格入库、反馈看板、Token 用量和运行状态。"],
            ["外观与退出", "顶部深色模式和“退出”。", "可切换深色模式；测试结束后退出账号。"],
        ],
        [1.25, 1.9, 3.35],
    )

    doc.add_heading("三、全功能内测清单", level=1)
    add_table(
        doc,
        ["模块", "必须测试", "通过标准"],
        [
            ["登录与会话", "登录、退出、新建对话、恢复历史会话、删除会话。", "登录稳定；历史内容不串号；删除后列表刷新正常。"],
            ["智能客服问答", "NAS/网盘、场地收费、制度流程、身份能力说明。", "同义问法回答一致；关键地址、金额、日期等数据准确；来源相关。"],
            ["公文协作", "通知、报告、建议、请示或议案；至少 3 轮连续修改。", "能完成生成、审核、修订，不中途退回，不泄露思考过程，不编造事实。"],
            ["文档编辑器", "打开/返回、自动导入生成稿、手动编辑、改字体字号行距、插入来源、清空、复制。", "编辑区可用；字数、内容状态、来源数量、自动保存状态更新合理。"],
            ["Word 导出", "自动模板、普通公文、院务会议案；修改后再导出。", "docx 可打开；内容与编辑器一致；标题、正文、落款、日期和页边距无明显异常。"],
            ["附件上传", "团队知识库上传、当前对话临时上传、移除附件、清除附件。", "上传状态清晰；临时文件不进入团队库；权限不足时提示明确。"],
            ["Excel/CSV", "上传表格，提问数据内容；按规则筛选、排序并导出。", "数值、日期、行数与原表一致；处理结果可下载并打开。"],
            ["报销表模板", "要求导出差旅费、会议费、劳务费/专家咨询费、其他费用报销表。", "能识别具体模板；不明确时主动询问；下载文件名和模板正确。"],
            ["内测反馈", "提交一般反馈、问题故障、生成质量、知识库检索、上传入库、导出 Word、响应速度。", "反馈可提交；后台能看到对应记录；分类和评分保存正确。"],
            ["后台与日志", "管理员查看知识库、审计日志、表格、反馈看板、Token 用量、服务健康。", "上传删除记录一致；监控无异常；服务健康接口正常。"],
        ],
        [1.3, 2.65, 2.55],
    )

    doc.add_heading("四、推荐测试范围", level=1)
    add_table(
        doc,
        ["场景", "建议提问或操作", "观察重点"],
        [
            ["NAS/网盘", "网盘地址是什么？如何使用 NAS 服务器？", "是否稳定输出 \\\\172.16.12.126，以及 Windows + R 操作步骤。"],
            ["场地收费", "场地使用收费表内容是什么？教室和会议室怎么收费？", "是否命中新版收费表，金额、人数、门牌号是否与 Excel 一致。"],
            ["制度流程", "某项制度怎么查？某流程如何办理？", "答案是否引用相关文件，是否避免无依据推测。"],
            ["公文生成", "帮我写一份通知/报告/对策建议。", "是否生成完整正文，是否避免脑补地点门牌、讲师、附件、调休规则等。"],
            ["多 Agent 协作", "生成公文后继续要求修改、润色、补充依据或压缩篇幅。", "规划、撰写、审核、修订是否衔接稳定，是否中途退回或泄露思考过程。"],
            ["文档编辑器", "打开文档编辑器，编辑生成稿，插入来源，改字体字号行距，复制内容。", "自动导入、手动编辑、自动保存、来源区、复制和清空是否正常。"],
            ["Word 导出", "公文生成完成后点击导出 Word，并打开导出的 docx 核对。", "标题、正文、落款、日期、页边距、字体字号和模板选择是否符合预期。"],
            ["Excel 与表格", "上传表格后提问、筛选、排序并导出。", "数据是否准确，导出的 Excel 是否可打开，行数和字段是否正确。"],
            ["反馈与后台", "提交反馈，管理员查看反馈看板、知识库审计和表格入库。", "前台提交和后台记录是否一致，是否便于定位问题。"],
            ["知识库管理", "上传、删除、重建索引。", "后台状态、审计日志、监控输出是否一致。"],
        ],
        [1.25, 2.45, 2.8],
    )

    doc.add_heading("五、基础使用步骤", level=1)
    add_numbered(doc, "登录系统后进入“智能体对话”页面。")
    add_numbered(doc, "输入问题或任务。知识库问答请尽量说明文件、业务主题或要查询的数据字段。")
    add_numbered(doc, "如果是公文生成，请给出文种、主题、时间、地点、对象、主要内容和落款单位。")
    add_numbered(doc, "查看回答中的来源文件和关键数据。涉及 Excel 金额、人数、日期时请人工抽查原表。")
    add_numbered(doc, "公文生成完成后，检查右侧编辑器内容，再选择 Word 模板并导出 docx。")
    add_numbered(doc, "发现错答、漏答、引用错误或导出版式异常时，按反馈格式记录并提交给管理员。")

    doc.add_heading("六、公文生成多 Agent 协作专项测试", level=1)
    add_callout(
        doc,
        "测试目的",
        "重点验证公文生成链路是否能稳定完成“需求理解、结构规划、正文撰写、事实边界审核、按反馈修订”。"
        "内测时不要只看文采，还要看是否完整、可控、不中断、不编造。",
    )
    add_table(
        doc,
        ["测试项", "建议输入", "通过标准"],
        [
            ["一次成稿", "帮我写一份关于开展信息化系统内测的通知，面向各科室，要求下周五前反馈问题。", "能输出完整通知，包含标题、主送对象、正文、要求、落款和日期，不出现无依据地点或附件。"],
            ["连续修订", "把刚才的通知改得更正式，并增加“问题反馈格式”要求。", "能继承上文，不重启话题；修改后结构更清晰，新增要求合理。"],
            ["压缩篇幅", "压缩到 300 字以内，保留关键要求。", "能保留核心事项、时间节点和反馈要求，不遗漏关键约束。"],
            ["事实边界", "生成一份会议通知，但我没有提供地点和主讲人。", "应主动使用待定或请补充，不编造教室、门牌、人员、附件。"],
            ["审核修订", "检查这份通知有没有不严谨的地方，并直接给出修订稿。", "能指出风险并给出修订结果，不暴露内部推理或 agent 调度细节。"],
            ["流式输出", "观察生成过程。", "页面应边生成边滚动，不能生成一半退回，不能先输出思考过程再改写。"],
        ],
        [1.2, 2.75, 2.55],
    )
    add_bullet(doc, "同一任务建议至少测 3 轮连续追问，确认多 Agent 状态不会丢失。")
    add_bullet(doc, "对正式公文，重点检查标题、主送、正文层次、落款、日期和事实边界。")
    add_bullet(doc, "如果出现“生成一半退回”“证据不足误判”“泄露分析过程”“凭空补充事实”，按高优先级反馈。")

    doc.add_heading("七、公文生成与导出 Word 专项测试", level=1)
    add_callout(
        doc,
        "测试目的",
        "重点验证生成内容是否能正确进入右侧文档编辑器，并通过“导出 Word”生成可打开、可编辑、版式合理的 docx 文件。"
        "内测人员应同时核对页面内容和导出文件，不能只看聊天回答。",
    )
    add_table(
        doc,
        ["测试项", "建议操作", "通过标准"],
        [
            ["普通通知导出", "生成一份通知，包含时间、地点、参加人员、事项要求、落款和日期，然后点击“导出 Word”。", "能下载 docx；打开后标题居中、正文清晰、落款和日期位置合理。"],
            ["模板自动识别", "保持 Word 模板为“自动识别”，分别导出通知、报告、议案类内容。", "普通公文走普通模板；议案类内容能识别为院务会议案模板或由用户手动选择。"],
            ["编辑器一致性", "导出前先核对右侧编辑器内容，再打开 docx 对照。", "导出文件内容与编辑器主体一致，不缺段、不重复、不混入上一轮对话。"],
            ["版式检查", "打开导出的 docx，检查首页、正文段落、编号、附件、落款、日期。", "无文字重叠、无异常空白页、无表格/段落被压坏，中文字体和行距可接受。"],
            ["修改后再导出", "先要求系统把通知改正式或压缩篇幅，再重新导出 Word。", "导出的是最新修订稿，不是上一版内容。"],
            ["异常处理", "用空编辑器、超短内容或含特殊符号内容尝试导出。", "空内容应提示不能导出；特殊符号不应导致 500 错误或文件名异常。"],
        ],
        [1.25, 2.75, 2.5],
    )
    add_bullet(doc, "每轮公文内测至少保留一个导出的 docx 文件，方便管理员复核版式和内容一致性。")
    add_bullet(doc, "如果导出的 Word 与页面显示不一致，应同时反馈原始提问、页面截图、导出的 docx 文件和操作时间。")
    add_bullet(doc, "导出文件用于内测核验，不代表正式发文；正式对外前仍需人工校对和审批。")

    doc.add_heading("八、内测注意事项", level=1)
    add_callout(
        doc,
        "数据边界",
        "暂不建议上传高度敏感的人事隐私、未公开合同全文、大量财务明细、个人身份信息或需要严格法律/财务结论的材料。"
        "如确需测试敏感材料，请先确认模型调用范围和脱敏策略。",
        fill="FFF4E5",
    )
    add_bullet(doc, "问法不同但意思相同的问题，应观察回答是否一致，例如“网盘”“NAS”“存储服务器”。")
    add_bullet(doc, "Excel 表格类问题，应重点核对金额、日期、人数、门牌号等原始字段。")
    add_bullet(doc, "文档编辑器测试应覆盖打开/返回、自动导入、手动编辑、来源插入、复制、清空、自动保存和导出。")
    add_bullet(doc, "公文生成内容如果出现用户未提供的具体地点、讲师、附件、政策或数据，应记录为问题。")
    add_bullet(doc, "公文生成多 Agent 测试应覆盖一次成稿、连续修订、压缩篇幅、事实审核和流式输出。")
    add_bullet(doc, "公文导出测试应覆盖自动模板、手动模板、修改后再导出、打开 docx 版式核验。")
    add_bullet(doc, "引用来源应与答案内容相关；如果混入无关文件，请截图并记录原问题。")
    add_bullet(doc, "不要把系统回答直接作为正式对外文件发布，内测阶段仍需人工复核。")

    doc.add_heading("九、问题反馈格式", level=1)
    add_table(
        doc,
        ["字段", "填写说明"],
        [
            ["问题类型", "登录会话 / 知识库问答 / Excel 数据 / 公文生成 / 多 Agent 协作 / 文档编辑器 / Word 导出 / 上传删除 / 后台管理 / 页面交互 / 其他"],
            ["原始提问", "完整复制用户当时输入的问题。"],
            ["实际回答", "粘贴回答关键片段，或附页面截图；导出问题请同时附 docx 文件。"],
            ["期望回答", "说明正确答案或期望行为。"],
            ["涉及文件", "如有，请写明文件名、表格名、页码或行号。"],
            ["是否可复现", "可复现 / 偶发 / 暂不确定。"],
            ["备注", "浏览器、时间、账号、操作步骤等。"],
        ],
        [1.35, 5.15],
    )

    doc.add_heading("十、管理员检查项", level=1)
    add_table(
        doc,
        ["检查项", "命令或位置", "正常结果"],
        [
            ["服务健康", "curl http://127.0.0.1:5003/api/health", "status 为 ok，storage.ok 为 true。"],
            ["进程状态", "pgrep -fl \"python.*app.py --port 5003|monitor_knowledge_ops.py\"", "同时看到 Web 服务和知识库监控进程。"],
            ["关键回归测试", ".venv311/bin/python -m pytest test_chat_rag.py test_rag_center_files.py test_chat_draft.py test_writer_agent.py test_spreadsheet_upload.py -q", "当前基线为 59 passed, 8 skipped。"],
            ["公文生成链路", "聊天页面、公文生成日志、document_stream_runner 输出", "无 reasoning_chunk 外泄，无中途退回，writer/reviewer 约束生效。"],
            ["文档编辑器链路", "浏览器页面、localStorage、编辑器状态面板", "生成稿自动导入；编辑、复制、来源插入、清空和自动保存正常。"],
            ["Word 导出链路", "聊天页面导出 Word、/api/export_docx、导出的 docx 文件", "接口返回 docx；文件可打开；内容与编辑器一致；版式无明显异常。"],
            ["知识库操作", "后台知识库文件、审计日志、表格入库页面", "上传、删除、归档、重建记录一致。"],
        ],
        [1.25, 3.65, 1.6],
    )

    doc.add_heading("十一、内测开放建议", level=1)
    add_bullet(doc, "第一阶段控制在 3-5 人，连续测试 3-5 个工作日。")
    add_bullet(doc, "每位用户每天至少提交 10 个真实办公问题，覆盖问答、公文生成、文档编辑器、Word/Excel 导出和文件管理。")
    add_bullet(doc, "管理员每日汇总问题清单，优先修复高频问法不一致、Excel 数字错误、来源混入、事实脑补和导出版式异常。")
    add_bullet(doc, "稳定后再扩大到 10-15 人，不建议一开始全员开放。")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
    print(OUTPUT)
