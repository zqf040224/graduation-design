"""
文档解析模块 - 支持 PDF、Word、TXT 格式解析

支持两种解析模式：
1. 纯文本模式（parse_document）：仅提取文字内容
2. 带格式模式（parse_document_with_format）：提取文字 + 格式信息
"""

from pathlib import Path
from typing import Optional, Dict, List
from dataclasses import dataclass
import json


@dataclass
class FormatFingerprint:
    """格式指纹 - 描述一段文字的格式特征"""
    font: str = ""           # 字体（如"黑体"、"仿宋_GB2312"）
    size: str = ""           # 字号（如"二号"、"三号"）
    bold: bool = False       # 是否加粗
    italic: bool = False     # 是否斜体
    alignment: str = ""      # 对齐方式（"居中"、"居左"、"居右"）
    first_line_indent: int = 0  # 首行缩进（字符数）
    line_spacing: str = ""   # 行距（如"28磅"）

    def to_dict(self) -> dict:
        return {
            "font": self.font,
            "size": self.size,
            "bold": self.bold,
            "italic": self.italic,
            "alignment": self.alignment,
            "first_line_indent": self.first_line_indent,
            "line_spacing": self.line_spacing
        }

    @classmethod
    def from_run(cls, run, para) -> "FormatFingerprint":
        """从 python-docx 的 run 和 para 对象提取格式"""
        fp = cls()

        # 字体
        east_asia_font = ""
        if run._element.rPr is not None and run._element.rPr.rFonts is not None:
            east_asia_font = (
                run._element.rPr.rFonts.get(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia"
                )
                or ""
            )
        if east_asia_font:
            fp.font = east_asia_font
        elif run.font.name:
            fp.font = run.font.name

        # 字号（转换为"一号"、"二号"等）
        if run.font.size:
            from docx.shared import Pt
            size_pt = run.font.size.pt if hasattr(run.font.size, 'pt') else run.font.size / 12700
            fp.size = pt_to_chinese_size(size_pt)

        # 加粗
        fp.bold = run.font.bold == True

        # 斜体
        fp.italic = run.font.italic == True

        # 对齐方式
        alignment_map = {
            0: "居左",
            1: "居中",
            2: "居右",
            3: "两端对齐",
        }
        if para.alignment is not None:
            fp.alignment = alignment_map.get(int(para.alignment), "")

        # 首行缩进
        if para.paragraph_format.first_line_indent:
            # 转换为字符数（每个字符约 0.35cm）
            indent_cm = para.paragraph_format.first_line_indent.cm
            fp.first_line_indent = int(indent_cm / 0.35)

        # 行距
        if para.paragraph_format.line_spacing:
            line_spacing = para.paragraph_format.line_spacing
            if hasattr(line_spacing, "pt"):
                fp.line_spacing = f"{int(round(line_spacing.pt))}磅"
            elif isinstance(line_spacing, (int, float)) and line_spacing > 1000:
                fp.line_spacing = f"{int(round(line_spacing / 12700))}磅"
            else:
                fp.line_spacing = str(line_spacing)

        return fp


def pt_to_chinese_size(pt: float) -> str:
    """将磅值转换为中文字号"""
    size_map = {
        26: "一号",
        22: "二号",
        16: "三号",
        14: "小四号",
        12: "四号",
        10.5: "五号",
        9: "小五号",
        8: "六号",
        7.5: "小六号",
        6.5: "七号",
        5.5: "八号"
    }
    return size_map.get(round(pt), f"{pt}磅")


def chinese_size_to_pt(size: str) -> float:
    """将中文字号转换为磅值"""
    size_map = {
        "一号": 22,
        "二号": 18,
        "三号": 16,
        "四号": 14,
        "小四号": 12,
        "五号": 10.5,
        "小五号": 9,
        "六号": 8,
        "小六号": 7.5,
        "七号": 6.5,
        "八号": 5.5
    }
    return size_map.get(size, 16)  # 默认三号（16磅）


def parse_document(file_path: Path) -> Optional[str]:
    """
    解析文档内容（纯文本模式）

    Args:
        file_path: 文件路径

    Returns:
        文档内容文本
    """
    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        return parse_pdf(file_path)
    elif suffix in [".docx", ".doc"]:
        return parse_docx(file_path)
    elif suffix == ".txt":
        return parse_txt(file_path)
    elif suffix == ".md":
        return parse_txt(file_path)
    else:
        print(f"⚠️ 不支持的文件格式：{suffix}")
        return None


def parse_document_with_format(file_path: Path) -> List[Dict]:
    """
    解析文档内容（带格式模式）

    Args:
        file_path: 文件路径

    Returns:
        包含文字和格式信息的列表
        [{
            "content": "文字内容",
            "format": FormatFingerprint
        }, ...]
    """
    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        return parse_pdf_with_format(file_path)
    elif suffix in [".docx", ".doc"]:
        return parse_docx_with_format(file_path)
    elif suffix == ".txt":
        return parse_txt_with_format(file_path)
    elif suffix == ".md":
        return parse_txt_with_format(file_path)
    else:
        print(f"⚠️ 不支持的文件格式：{suffix}")
        return []


def parse_pdf(file_path: Path) -> Optional[str]:
    """解析 PDF 文件"""
    pages = parse_pdf_pages(file_path)
    if not pages:
        return None
    return "\n".join(page["content"] for page in pages if page.get("content"))


def parse_pdf_pages(file_path: Path) -> List[Dict]:
    """按页解析 PDF，保留页码与轻量解析警告。"""
    try:
        from PyPDF2 import PdfReader

        reader = PdfReader(str(file_path))
        pages = []
        empty_pages = 0

        for idx, page in enumerate(reader.pages, start=1):
            text = page.extract_text()
            text = (text or "").strip()
            if text:
                pages.append({
                    "content": text,
                    "page_start": idx,
                    "page_end": idx,
                    "format": FormatFingerprint().to_dict(),
                    "parse_warnings": [],
                })
            else:
                empty_pages += 1

        if empty_pages and pages:
            pages[0]["parse_warnings"] = [
                f"PDF 中有 {empty_pages} 页未提取到文本，可能是扫描页或图片页"
            ]
        return pages

    except Exception as e:
        print(f"  ❌ PDF 解析失败：{e}")
        return []


def parse_pdf_with_format(file_path: Path) -> List[Dict]:
    """
    解析 PDF 文件（带格式模式）

    注意：PDF 格式信息较少，主要返回纯文本
    """
    return parse_pdf_pages(file_path)


def parse_docx(file_path: Path) -> Optional[str]:
    """解析 Word 文件（纯文本）"""
    try:
        from docx import Document

        doc = Document(str(file_path))
        content = []

        # 提取段落
        for para in doc.paragraphs:
            if para.text.strip():
                content.append(para.text)

        # 提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text)
                if row_text:
                    content.append(" | ".join(row_text))

        return "\n".join(content)

    except Exception as e:
        print(f"  ❌ Word 解析失败：{e}")
        return None


def parse_docx_with_format(file_path: Path) -> List[Dict]:
    """
    解析 Word 文件（带格式模式）

    提取每个段落的文字和格式信息
    """
    try:
        from docx import Document

        doc = Document(str(file_path))
        results = []

        # 处理段落
        for para in doc.paragraphs:
            if not para.text.strip():
                continue

            # 如果段落内有多个 run（不同格式），需要拆分
            if len(para.runs) > 1:
                for run in para.runs:
                    if run.text.strip():
                        fp = FormatFingerprint.from_run(run, para)
                        results.append({
                            "content": run.text,
                            "format": fp.to_dict()
                        })
            else:
                # 整个段落同一格式
                run = para.runs[0] if para.runs else None
                if run:
                    fp = FormatFingerprint.from_run(run, para)
                else:
                    fp = FormatFingerprint()
                results.append({
                    "content": para.text,
                    "format": fp.to_dict()
                })

        # 处理表格
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text)
                if row_text:
                    results.append({
                        "content": " | ".join(row_text),
                        "format": FormatFingerprint().to_dict()
                    })

        return results

    except Exception as e:
        print(f"  ❌ Word 解析失败：{e}")
        return []


def parse_txt(file_path: Path) -> Optional[str]:
    """解析 TXT/MD 文件"""
    try:
        # 尝试不同编码
        encodings = ["utf-8", "gbk", "gb2312", "utf-8-sig"]

        for encoding in encodings:
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue

        print(f"  ❌ 无法识别文件编码")
        return None

    except Exception as e:
        print(f"  ❌ TXT 解析失败：{e}")
        return None


def parse_txt_with_format(file_path: Path) -> List[Dict]:
    """
    解析 TXT/MD 文件（带格式模式）

    TXT 文件无格式信息，返回默认格式
    """
    text = parse_txt(file_path)
    if not text:
        return []

    default_format = FormatFingerprint()
    return [{
        "content": text,
        "format": default_format.to_dict()
    }]


def format_to_style_spec(format_dict: dict) -> str:
    """
    将格式指纹转换为 Prompt 中的样式描述

    用于生成 Prompt 时描述格式要求
    """
    specs = []

    if format_dict.get("font"):
        specs.append(f"字体：{format_dict['font']}")

    if format_dict.get("size"):
        specs.append(f"字号：{format_dict['size']}")

    if format_dict.get("bold"):
        specs.append("加粗")

    if format_dict.get("alignment"):
        specs.append(f"对齐：{format_dict['alignment']}")

    if format_dict.get("line_spacing"):
        specs.append(f"行距：{format_dict['line_spacing']}")

    if format_dict.get("first_line_indent"):
        specs.append(f"首行缩进：{format_dict['first_line_indent']}字符")

    return "，".join(specs) if specs else "默认格式"


if __name__ == "__main__":
    # 测试
    import sys

    if len(sys.argv) > 1:
        file_path = Path(sys.argv[1])

        print("=" * 60)
        print("测试纯文本模式：")
        print("=" * 60)
        content = parse_document(file_path)
        if content:
            print(f"✓ 解析成功，内容长度：{len(content)} 字符")
            print(f"\n前 500 字符:\n{content[:500]}")
        else:
            print("解析失败")

        print("\n" + "=" * 60)
        print("测试带格式模式：")
        print("=" * 60)
        formatted = parse_document_with_format(file_path)
        if formatted:
            print(f"✓ 解析成功，共 {len(formatted)} 个片段")
            for i, item in enumerate(formatted[:5]):
                fmt = item["format"]
                style = format_to_style_spec(fmt)
                print(f"\n[{i+1}] {style}")
                print(f"    内容：{item['content'][:50]}...")
        else:
            print("解析失败")
